import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re

print("Synchronizing all Word documents to peak Q1 formatting with 6 embedded figures & superscript parsing...")

md_path = 'Manuscript_Student_Burnout.md'

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

doc = docx.Document()

# Set standard margins (1 inch on all sides)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_formatted_text(paragraph, text, base_bold=False, base_color=None, base_size=None):
    """
    Parses Markdown and HTML inline formatting:
    - <sup>...</sup> -> Superscript
    - <sub>...</sub> -> Subscript
    - **...**       -> Bold
    - *...*         -> Italic
    """
    pattern = re.compile(r'(<sup>.*?</sup>|<sub>.*?</sub>|\*\*.*?\*\*|\*.*?\*)')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('<sup>') and part.endswith('</sup>'):
            content = part[5:-6]
            r = paragraph.add_run(content)
            r.font.superscript = True
            if base_bold: r.font.bold = True
            if base_color: r.font.color.rgb = base_color
            if base_size: r.font.size = base_size
        elif part.startswith('<sub>') and part.endswith('</sub>'):
            content = part[5:-6]
            r = paragraph.add_run(content)
            r.font.subscript = True
            if base_bold: r.font.bold = True
            if base_color: r.font.color.rgb = base_color
            if base_size: r.font.size = base_size
        elif part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            r = paragraph.add_run(content)
            r.font.bold = True
            if base_color: r.font.color.rgb = base_color
            if base_size: r.font.size = base_size
        elif part.startswith('*') and part.endswith('*'):
            content = part[1:-1]
            r = paragraph.add_run(content)
            r.font.italic = True
            if base_bold: r.font.bold = True
            if base_color: r.font.color.rgb = base_color
            if base_size: r.font.size = base_size
        else:
            r = paragraph.add_run(part)
            if base_bold: r.font.bold = True
            if base_color: r.font.color.rgb = base_color
            if base_size: r.font.size = base_size

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "color", "sz", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_table_styling(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell, 
                             top={"val": "single", "sz": "4", "color": "94A3B8"},
                             bottom={"val": "single", "sz": "4", "color": "94A3B8"},
                             left={"val": "none"}, right={"val": "none"})
            if i == 0: # Header row
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), '1E3A8A') # Deep Blue Header
                cell._tc.get_or_add_tcPr().append(shading)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    runs = list(p.runs)
                    p.clear()
                    for r in runs:
                        new_r = p.add_run(r.text)
                        new_r.font.bold = True
                        new_r.font.color.rgb = RGBColor(255, 255, 255)

# Process Markdown lines
i = 0
n = len(lines)

while i < n:
    line = lines[i].rstrip('\n')
    
    if line.strip() == '---' or line.strip() == '***':
        i += 1
        continue
    
    # Headings
    if line.startswith('# '):
        h = doc.add_heading(level=1)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        add_formatted_text(h, line[2:].strip(), base_color=RGBColor(15, 23, 42))
        i += 1
        continue
    elif line.startswith('## '):
        h = doc.add_heading(level=2)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        add_formatted_text(h, line[3:].strip(), base_color=RGBColor(30, 58, 138))
        i += 1
        continue
    elif line.startswith('### '):
        h = doc.add_heading(level=3)
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(3)
        add_formatted_text(h, line[4:].strip(), base_color=RGBColor(4, 120, 87))
        i += 1
        continue
    elif line.startswith('#### '):
        h = doc.add_heading(level=4)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        add_formatted_text(h, line[5:].strip())
        i += 1
        continue
        
    # Check for Markdown Image embed
    img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
    if img_match:
        caption, img_file = img_match.groups()
        img_path = img_file.strip()
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(12)
            p_img.paragraph_format.space_after = Pt(4)
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=Inches(5.8))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(14)
            add_formatted_text(p_cap, caption, base_color=RGBColor(71, 85, 105), base_size=Pt(9.5))
        i += 1
        continue

    # Check for Markdown Table
    if line.strip().startswith('|') and '|' in line[1:]:
        table_lines = []
        while i < n and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        
        rows = []
        for r in table_lines:
            cells = [c.strip() for c in r.strip('|').split('|')]
            if all(re.match(r'^:?-+:?$', c) for c in cells if c):
                continue
            rows.append(r)
        
        if rows:
            parsed_rows = [[c.strip() for c in r.strip('|').split('|')] for r in rows]
            num_rows = len(parsed_rows)
            num_cols = max(len(r) for r in parsed_rows)
            
            table = doc.add_table(rows=num_rows, cols=num_cols)
            for r_idx, r_data in enumerate(parsed_rows):
                for c_idx, val in enumerate(r_data):
                    if c_idx < num_cols:
                        cell = table.cell(r_idx, c_idx)
                        p_cell = cell.paragraphs[0]
                        p_cell.text = "" # Clear default text
                        add_formatted_text(p_cell, val)
            set_table_styling(table)
            
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(0)
            p_space.paragraph_format.space_after = Pt(6)
        continue

    # Regular Paragraph / Bullet Points
    line_str = line.strip()
    if line_str:
        if line_str.startswith('- ') or line_str.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text_content = line_str[2:].strip()
        else:
            p = doc.add_paragraph()
            text_content = line_str
        
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        add_formatted_text(p, text_content)
    i += 1

output_file = 'Manuscript_Student_Burnout.docx'

try:
    doc.save(output_file)
    print(f"Word document synchronized: '{output_file}'")
except PermissionError:
    print(f"Main file '{output_file}' is open in Word.")

